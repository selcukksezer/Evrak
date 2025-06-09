# functions/bep_plan_table_generator.py
import json
import io
import datetime
from firebase_functions import https_fn, options
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# --- YARDIMCI FONKSİYONLAR ---

def set_cell_shading(cell, fill_color_hex_string): # Örn: "4F81BD"
    """Hücrenin arkaplan rengini ayarlar."""
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{fill_color_hex_string}" w:val="clear" />'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_text_formatted(cell, text, bold=False, align_center=False, font_size_pt=None, font_color_rgb_obj=None):
    """Hücreye metin ekler ve formatlama yapar. Paragraphs[0] kullanılır."""
    if not cell.paragraphs:
        cell.add_paragraph()
    p = cell.paragraphs[0]
    p.clear()  # Mevcut içeriği ve formatlamayı temizle
    run = p.add_run(str(text) if text is not None else '')
    run.bold = bold
    if font_size_pt:
        run.font.size = Pt(font_size_pt)
    if font_color_rgb_obj:
        run.font.color.rgb = font_color_rgb_obj  # RGBColor nesnesi bekler

    if align_center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

# --- BEP PLAN TABLOSU OLUŞTURMA FONKSİYONU ---
def create_bep_table_only(document, data):
    """Sadece Bireyselleştirilmiş Eğitim Planı tablosunu oluşturur."""
    document.add_heading('III-BİREYSELLEŞTİRİLMİŞ EĞİTİM PLANI', level=1)

    if not isinstance(data, dict):
        print("Hata: Ana veri bir sözlük değil.")
        document.add_paragraph("Hatalı veri formatı.")
        return

    dersler_raw = data.get("secilenDersler")
    dersler = dersler_raw if isinstance(dersler_raw, list) else []

    if not dersler:
        document.add_paragraph("Öğrenci için seçilmiş ders veya hedef bulunmamaktadır.")
        return

    plan_table = document.add_table(rows=2, cols=9)
    plan_table.style = 'Table Grid' # Temel tablo stili

    # Sabit renkler ve başlıklar
    headers1 = ["Uzun Dönemli Amaçlar", "Kısa Dönemli Amaçlar", "Ölçüt", "Yöntem ve Teknik", "Kullanılacak Materyaller", "Başlama ve Bitiş Tarihi", "Ölçme ve Değerlendirme", "", ""]
    headers2 = ["", "", "", "", "", "", "Değerlendirme Yöntem ve Teknikleri", "Değerlendirme Tarihleri", "Performans"]

    header_font_color = RGBColor(0xFF, 0xFF, 0xFF) # Beyaz
    header_bg_color = '4F81BD' # Koyu Mavi
    uda_ders_bg_color = 'D3DFEE' # Açık Mavi

    # Başlık satırlarını doldur ve formatla
    for i, text in enumerate(headers1):
        cell = plan_table.cell(0, i)
        set_cell_text_formatted(cell, text, bold=True, font_color_rgb_obj=header_font_color, align_center=True)
        set_cell_shading(cell, header_bg_color)

    for i, text in enumerate(headers2):
        cell = plan_table.cell(1, i)
        set_cell_text_formatted(cell, text, bold=True, font_color_rgb_obj=header_font_color, align_center=True)
        set_cell_shading(cell, header_bg_color)

    # Başlık hücrelerini birleştir
    for i in range(6): # 0-5 arası sütunlar
        plan_table.cell(0, i).merge(plan_table.cell(1, i))
    plan_table.cell(0, 6).merge(plan_table.cell(0, 8)) # "Ölçme ve Değerlendirme" başlığı

    # Veri satırlarını doldur
    for ders_item in dersler:
        if not isinstance(ders_item, dict):
            print(f"Uyarı: 'secilenDersler' içinde sözlük olmayan öğe: {ders_item}")
            continue

        # Ders adı satırı
        ders_row_cells = plan_table.add_row().cells
        ders_cell = ders_row_cells[0]
        ders_cell.merge(ders_row_cells[8]) # Tüm sütunları birleştir
        set_cell_text_formatted(ders_cell, ders_item.get("dersAdi", "").upper(), bold=True, align_center=True)
        set_cell_shading(ders_cell, uda_ders_bg_color)

        uzun_donemli_amaclar_raw = ders_item.get("uzunDonemliAmaclar")
        uzun_donemli_amaclar = uzun_donemli_amaclar_raw if isinstance(uzun_donemli_amaclar_raw, list) else []

        for uda_item in uzun_donemli_amaclar:
            if not isinstance(uda_item, dict):
                print(f"Uyarı: 'uzunDonemliAmaclar' içinde sözlük olmayan öğe: {uda_item}")
                continue

            kda_list_raw = uda_item.get("kisaDonemliAmaclar")
            kda_list = kda_list_raw if isinstance(kda_list_raw, list) else []

            if not kda_list:
                continue

            num_kdas_for_this_uda = len(kda_list)
            first_kda_row_for_this_uda_idx = -1

            for kda_idx, kda_item in enumerate(kda_list):
                if not isinstance(kda_item, dict):
                    print(f"Uyarı: 'kisaDonemliAmaclar' içinde sözlük olmayan öğe: {kda_item}")
                    continue

                kda_row_cells = plan_table.add_row().cells
                current_row_idx = len(plan_table.rows) - 1

                if kda_idx == 0:
                    first_kda_row_for_this_uda_idx = current_row_idx
                    uda_cell = kda_row_cells[0]
                    set_cell_text_formatted(uda_cell, uda_item.get("udaMetni", ""), align_center=True)
                    set_cell_shading(uda_cell, uda_ders_bg_color)
                else:
                    set_cell_shading(kda_row_cells[0], uda_ders_bg_color)

                kda_row_cells[1].text = str(kda_item.get("kdaMetni") or "Yok")
                kda_row_cells[2].text = str(kda_item.get("olcut") or "Yok")

                ogretim_yontemleri_raw = kda_item.get("ogretimYontemleri")
                ogretim_yontemleri = ogretim_yontemleri_raw if isinstance(ogretim_yontemleri_raw, list) else []
                kda_row_cells[3].text = ", ".join(map(str, ogretim_yontemleri)) if ogretim_yontemleri else "Yok"

                kullanilan_materyaller_raw = kda_item.get("kullanilanMateryaller")
                kullanilan_materyaller = kullanilan_materyaller_raw if isinstance(kullanilan_materyaller_raw, list) else []
                kda_row_cells[4].text = ", ".join(map(str, kullanilan_materyaller)) if kullanilan_materyaller else "Yok"

                baslama_tarihi_val = kda_item.get('baslamaTarihi')
                bitis_tarihi_val = kda_item.get('bitisTarihi')
                baslama_tarihi_str = str(baslama_tarihi_val or "Yok")
                bitis_tarihi_str = str(bitis_tarihi_val or "Yok")
                kda_row_cells[5].text = f"{baslama_tarihi_str} - {bitis_tarihi_str}"

                kda_row_cells[6].text = "Gözlem, Kontrol Listesi" # Bu sabit kalabilir veya veriden alınabilir

                degerlendirme_tarihi_val = kda_item.get("degerlendirmeTarihi")
                # Eğer değerlendirme tarihi yoksa bitiş tarihini kullan, o da yoksa "Yok" yaz
                final_degerlendirme_tarihi_str = str((degerlendirme_tarihi_val or bitis_tarihi_val) or "Yok")
                kda_row_cells[7].text = final_degerlendirme_tarihi_str
                kda_row_cells[8].text = "" # Performans alanı boş

                for cell_idx in range(1, 9):
                    kda_row_cells[cell_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            if num_kdas_for_this_uda > 0 and first_kda_row_for_this_uda_idx != -1:
                if num_kdas_for_this_uda > 1:
                    cell_to_merge_from = plan_table.cell(first_kda_row_for_this_uda_idx, 0)
                    cell_to_merge_to = plan_table.cell(len(plan_table.rows) - 1, 0)
                    cell_to_merge_from.merge(cell_to_merge_to)

                plan_table.cell(first_kda_row_for_this_uda_idx, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

# --- ANA CLOUD FUNCTION ---
@https_fn.on_request(memory=options.MemoryOption.MB_256, timeout_sec=120)
def generate_bep_plan_table_only(req: https_fn.Request) -> https_fn.Response:
    """Flutter'dan gelen BEP Planı verileriyle SADECE BEP PLAN TABLOSUNU içeren Word belgesi oluşturur."""
    try:
        data = req.get_json(silent=True) # HTTP POST isteği için JSON gövdesini al
        if not data:
            raise ValueError("Gelen veri JSON formatında değil veya boş.")

        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        unique_filename = f"BEP_Plani_Tablosu_{timestamp}.docx"

        document = Document()

        section = document.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height

        create_bep_table_only(document, data)

        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)

        return https_fn.Response(
            buffer.getvalue(),
            status=200,
            headers={
                "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "Content-Disposition": f"attachment; filename=\"{unique_filename}\""
            }
        )
    except Exception as e:
        print(f"Hata oluştu (generate_bep_plan_table_only): {e}")
        error_payload = json.dumps({"error": f"Sunucuda bir hata oluştu: {str(e)}"})
        return https_fn.Response(error_payload, status=500, mimetype="application/json")
