from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import io

# Helper function to set cell background color
def set_cell_shading(cell, fill_color_hex_string): # e.g., "A9D18E" (açık yeşil)
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{fill_color_hex_string}" />'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)

def generate_kaba_degerlendirme_docx(data):
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    section = document.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.7) # 7.1 inç kullanılabilir genişlik için
    section.right_margin = Inches(0.7)

    # Yeşil Tonları
    dark_green_hex = "548235" # Koyu Yeşil (Başlıklar için)
    light_green_hex = "A9D18E" # Açık Yeşil (Tablo başlık arkaplanı için)
    white_font_color = RGBColor(0xFF, 0xFF, 0xFF)

    # 1. Okul Adı
    school_name_paragraph = document.add_paragraph()
    school_name_run = school_name_paragraph.add_run(data.get('kurum_adi', 'Kurum Adı Girilmemiş').upper())
    school_name_run.font.size = Pt(14)
    school_name_run.font.bold = True
    school_name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 2. Form Başlığı
    form_title_paragraph = document.add_paragraph()
    form_title_run = form_title_paragraph.add_run('KABA DEĞERLENDİRME FORMU')
    form_title_run.font.size = Pt(13)
    form_title_run.font.bold = True
    # form_title_run.font.color.rgb = RGBColor.from_string(dark_green_hex) # İsteğe bağlı renk
    form_title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph() # Başlık ile öğrenci bilgileri arasına boşluk

    # 3. Öğrenci Bilgileri Tablosu
    student_info_table = document.add_table(rows=1, cols=2)
    student_info_table.style = 'Table Grid'
    student_info_table.autofit = False
    student_info_table.allow_autofit = False
    student_info_table.columns[0].width = Inches(2.0) # Etiket sütunu
    student_info_table.columns[1].width = Inches(5.1) # Değer sütunu (7.1 - 2.0)

    # Tablo Başlığı
    hdr_cell_student = student_info_table.cell(0, 0)
    hdr_cell_student.merge(student_info_table.cell(0, 1))
    p_student_hdr = hdr_cell_student.paragraphs[0]
    run_student_hdr = p_student_hdr.add_run('ÖĞRENCİ BİLGİLERİ')
    run_student_hdr.font.bold = True
    run_student_hdr.font.color.rgb = white_font_color
    p_student_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_shading(hdr_cell_student, dark_green_hex)

    # Öğrenci Bilgileri Satırları
    student_info_data = [
        ('Öğrenci Adı-Soyadı:', data.get('ogrenci_adi', '')),
        ('Uygulayıcı Adı-Soyadı:', data.get('uygulayici_adi', '')),
        ('Uygulama Tarihi:', data.get('uygulama_tarihi', '')),
    ]

    for label, value in student_info_data:
        cells = student_info_table.add_row().cells
        cells[0].paragraphs[0].add_run(label).bold = True
        cells[1].text = value

    document.add_paragraph() # Öğrenci bilgileri tablosu ile dersler arasına boşluk

    # 4. Dersler ve Değerlendirme Tabloları (Sayfa sonu yok)
    for ders_data in data.get('dersler', []):
        ders_heading_paragraph = document.add_paragraph()
        ders_heading_run = ders_heading_paragraph.add_run(f"DERSİN ADI: {ders_data.get('ders_adi', '').upper()}")
        ders_heading_run.bold = True
        ders_heading_run.font.size = Pt(12)
        # ders_heading_run.font.color.rgb = RGBColor.from_string(dark_green_hex) # İsteğe bağlı renk
        ders_heading_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # document.add_paragraph() # Ders başlığı ile tablo arası boşluk azaltıldı

        eval_table = document.add_table(rows=1, cols=3)
        eval_table.style = 'Table Grid'
        eval_table.autofit = False
        eval_table.allow_autofit = False
        eval_table.columns[0].width = Inches(5.8) # Hedef davranışlar (7.1 - 0.65 - 0.65)
        eval_table.columns[1].width = Inches(0.65) # Evet
        eval_table.columns[2].width = Inches(0.65) # Hayır

        hdr_cells_eval = eval_table.rows[0].cells
        hdr_cells_eval[0].text = 'Hedef Davranışlar'
        hdr_cells_eval[1].text = 'Evet'
        hdr_cells_eval[2].text = 'Hayır'
        for cell_idx, cell in enumerate(hdr_cells_eval):
            p_hdr_eval = cell.paragraphs[0]
            p_hdr_eval.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p_hdr_eval.runs:
                run.font.bold = True
                run.font.color.rgb = white_font_color
            set_cell_shading(cell, light_green_hex)
            if cell_idx == 0 : # Hedef davranışlar başlığı için koyu yeşil
                 set_cell_shading(cell, dark_green_hex)

        for uda_data in ders_data.get('uzun_donemli_amaclar', []):
            uda_row = eval_table.add_row()
            uda_cell = uda_row.cells[0]
            uda_cell.merge(uda_row.cells[2]) # UDA satırında hücreleri birleştir
            uda_p = uda_cell.paragraphs[0]
            uda_p.add_run(uda_data.get('kazanim_metni', '')).bold = True
            uda_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            for kda_data in uda_data.get('kisa_donemli_amaclar', []):
                kda_row_cells = eval_table.add_row().cells
                kda_row_cells[0].text = kda_data.get('kazanim_metni', '')
                mark_cell_idx = 1 if kda_data.get('basarili_mi', False) else 2
                kda_row_cells[mark_cell_idx].text = 'X'
                kda_row_cells[mark_cell_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        document.add_paragraph() # Ders tabloları arasına küçük bir boşluk

    # 5. İmza Bölümü (Son tablonun hemen altında)
    # document.add_paragraph() # Ekstra boşluk kaldırıldı, gerekirse eklenebilir
    p_imza = document.add_paragraph()
    p_imza.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_imza.add_run(f"{data.get('uygulayici_adi', '')}\n")
    p_imza.add_run("Ders Öğretmeni").bold = True

    doc_io = io.BytesIO()
    document.save(doc_io)
    doc_io.seek(0)
    return doc_io

dummy_data_for_word = {
    'kurum_adi': 'Milli Eğitim Bakanlığı Okulu',
    'ogrenci_adi': 'Ayşe Yıldız',
    'uygulayici_adi': 'Mehmet Öztürk',
    'uygulama_tarihi': '09.06.2025',
    'dersler': [
        {
            'ders_adi': 'Matematik',
            'uzun_donemli_amaclar': [
                {
                    'kazanim_metni': 'Sayıları tanır ve sıralar.',
                    'kisa_donemli_amaclar': [
                        {'kazanim_metni': '1\\\'den 20\\\'ye kadar ritmik sayar.', 'basarili_mi': True},
                        {'kazanim_metni': 'Verilen sayılar arasında büyüklük küçüklük ilişkisi kurar.', 'basarili_mi': False},
                        {'kazanim_metni': 'Nesneleri kullanarak toplama işlemi yapar.', 'basarili_mi': True},
                    ]
                },
                {
                    'kazanim_metni': 'Geometrik şekilleri tanır.',
                    'kisa_donemli_amaclar': [
                        {'kazanim_metni': 'Kare, üçgen ve daireyi ayırt eder.', 'basarili_mi': True},
                    ]
                }
            ]
        },
        {
            'ders_adi': 'Günlük Yaşam Becerileri',
            'uzun_donemli_amaclar': [
                {
                    'kazanim_metni': 'Kişisel bakımını yapar.',
                    'kisa_donemli_amaclar': [
                        {'kazanim_metni': 'Ellerini ve yüzünü yıkar.', 'basarili_mi': True},
                        {'kazanim_metni': 'Dişlerini fırçalar.', 'basarili_mi': False},
                    ]
                }
            ]
        }
    ]
}

if __name__ == '__main__':
    print("Word oluşturucu test modunda çalıştırılıyor...")
    doc_bytes_io = generate_kaba_degerlendirme_docx(dummy_data_for_word)
    try:
        with open('kaba_degerlendirme_test.docx', 'wb') as f:
            f.write(doc_bytes_io.read())
        print("Test belgesi \'kaba_degerlendirme_test.docx\' olarak başarıyla oluşturuldu.")
    except Exception as e:
        print(f"Test belgesi oluşturulurken hata: {e}")
