# main.py
import json
import io
import datetime
from firebase_functions import https_fn, options # options importunun varlığından emin olun
from firebase_admin import initialize_app
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.section import WD_ORIENT
# YENİ: Hücre arkaplan rengi için gerekli importlar
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# YENİ: Diğer dosyadan fonksiyonu içe aktar
from bep_plan_table_generator import generate_bep_plan_table_only # GÖRECELİ IMPORT DÜZELTİLDİ
from kaba_degerlendirme_word_generator import generate_kaba_degerlendirme_docx # YENİ KABA DEĞERLENDİRME İÇİN IMPORT

initialize_app()

# --- YENİ: STİL YARDIMCI FONKSİYONLARI ---

def set_cell_shading(cell, fill_color):
    """Hücrenin arkaplan rengini ayarlar."""
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{fill_color}" w:val="clear" />'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)

def apply_custom_table_style(table, header_color='#4F81BD', key_column_color='#D3DFEE'):
    """Tabloya özel mavi tonlarındaki stili uygular."""
    # Başlık satırını renklendir
    for cell in table.rows[0].cells:
        set_cell_shading(cell, header_color)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) # Beyaz yazı
                run.font.bold = True

    # Diğer satırları biçimlendir
    for i, row in enumerate(table.rows):
        if i == 0: continue # Başlık satırını atla
        # İlk sütunu (anahtar/etiket sütunu) açık mavi yap
        set_cell_shading(row.cells[0], key_column_color)
        # Diğer sütunlar (veri sütunları) beyaz kalacak

# --- BELGE OLUŞTURMA FONKSİYONLARI (GÜNCELLENDİ) ---

def create_cover_page(document, data):
    """Belgenin kapak sayfasını ve logosunu oluşturur."""
    # YENİ: Logo ekleme ve boşluk bırakma
    document.add_paragraph() # Üstten boşluk
    try:
        document.add_picture('meb_logo.png', width=Inches(2.0))
        # Eklenen son paragrafı (resmi içeren) ortala
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except FileNotFoundError:
        print("UYARI: meb_logo.png dosyası bulunamadı. Lütfen 'functions' klasörüne ekleyin.")
        document.add_paragraph("[MEB LOGOSU]").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Logo ile tablo arasında boşluk
    document.add_paragraph()
    document.add_paragraph()

    p_title = document.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.add_run('B��REYSELLEŞTİRİLMİŞ EĞİTİM PROGRAMI DOSYASI').bold = True
    document.add_paragraph()

    table = document.add_table(rows=1, cols=1) # Sadece başlık
    set_cell_text(table.cell(0, 0), 'ÖĞRENCİNİN', align_center=True, bold=True)

    # Dinamik verileri ekle
    row1 = table.add_row().cells
    set_cell_text(row1[0], f"ADI SOYADI: {data.get('ogrenciAdSoyad', '')}")
    row2 = table.add_row().cells
    set_cell_text(row2[0], f"OKULU: {data.get('calisilanOkul', '')}")
    row3 = table.add_row().cells
    set_cell_text(row3[0], f"NUMARASI: {data.get('ogrenciNumarasi', '')}")

    # YENİ: Stil uygulama
    apply_custom_table_style(table, header_color='#4F81BD', key_column_color='#FFFFFF') # Kapakta ilk sütun beyaz

# ... (Diğer yardımcı fonksiyonlar: set_cell_text vb. aynı kalabilir) ...
def set_cell_text(cell, text, bold=False, align_center=False):
    """Hücreye metin ekler ve temel formatlama yapar."""
    p = cell.paragraphs[0]
    p.text = str(text) if text is not None else ''
    p.runs[0].bold = bold
    if align_center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def create_student_info_table(document, data):
    """I-Öğrenci Bilgileri bölümünü oluşturur."""
    document.add_heading('I-Öğrenci Bilgileri', level=1)

    table_data = [
        ("Adı-Soyadı", data.get("ogrenciAdSoyad")),
        ("Sınıfı", f"{data.get('sinifDuzeyi', '')}/{data.get('subeAdi', '')}"),
        ("Okul numarası", data.get("ogrenciNumarasi")),
        ("Doğum tarihi", data.get("dogumTarihi")),
        ("İl/ilçe özel eğitim hizmetleri yerleştirme kurul kararı", data.get("kurulKarari")),
        ("Özel eğitim ihtiyacına yönelik aldığı eğitsel tanı", data.get("egitimselTani")),
        ("Varsa kullandığı destek materyalleri/cihazlar", ", ".join(data.get("kullanilanCihazlar", []))),
        ("Eğitim ortamına ilişkin düzenlemeler", data.get("egitimOrtamiDuzenlemesi")),
        ("BEP Başlangıç Tarihi", data.get("bepBaslangicTarihi")),
        ("BEP Bitiş Tarihi", data.get("bepBitisTarihi")),
    ]
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = 'Öğrenci Bilgileri'
    table.cell(0, 0).merge(table.cell(0, 1))

    for key, value in table_data:
        row_cells = table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = str(value) if value is not None else ''

    apply_custom_table_style(table)
    document.add_paragraph()

    document.add_heading('Aile ile ilgili bilgiler', level=2)
    family_table = document.add_table(rows=1, cols=4)
    headers = ['', 'Anne', 'Baba', 'Veli/Vasi']
    for i, header in enumerate(headers):
        family_table.cell(0, i).text = header

    veli_secimi = data.get("veliSecimi", "")
    veli_ad_soyad = ""
    if veli_secimi == "Anne": veli_ad_soyad = data.get("anneAdSoyad", "")
    elif veli_secimi == "Baba": veli_ad_soyad = data.get("babaAdSoyad", "")
    elif veli_secimi == "Diğer": veli_ad_soyad = data.get("digerVeliAdSoyad", "")

    rows_data = [
        ("Adı-Soyadı", [data.get("anneAdSoyad", ""), data.get("babaAdSoyad", ""), veli_ad_soyad]),
        ("Telefon", [data.get("anneTelefon", ""), data.get("babaTelefon", ""), ""])
    ]
    for key, values in rows_data:
        row_cells = family_table.add_row().cells
        row_cells[0].text = key
        for i, value in enumerate(values):
            row_cells[i+1].text = str(value) if value is not None else ''

    apply_custom_table_style(family_table)


def create_performance_form(document, data):
    """II-Eğitsel Performans Formu bölümünü oluşturur."""
    document.add_heading('II-Eğitsel Performans Formu', level=1)

    perf_table = document.add_table(rows=1, cols=2)
    perf_table.cell(0, 0).text = 'Gelişim alanları/Dersler'
    perf_table.cell(0, 1).text = 'Performans düzeyi'

    gelişim_oykusu_row = perf_table.add_row().cells
    gelişim_oykusu_row[0].text = "Öğrencinin Gelişim Öyküsü"
    gelişim_oykusu_row[0].merge(gelişim_oykusu_row[1])
    gelişim_oykusu_row[0].text = f"Öğrencinin Gelişim Öyküsü: {data.get('gelisimOykusu', '')}"


    dersler = data.get("secilenDersler", [])
    if dersler:
        for ders in dersler:
            row_cells = perf_table.add_row().cells
            row_cells[0].text = ders.get("dersAdi", "")
            perf_level_texts = [uda.get("udaMetni", "") for uda in ders.get("uzunDonemliAmaclar", [])]
            row_cells[1].text = " ".join(perf_level_texts)

    davranis_row = perf_table.add_row().cells
    davranis_row[0].text = "Varsa davranış problemini tanımlayınız"
    davranis_row[1].text = data.get("davranisProblemi", "")

    apply_custom_table_style(perf_table)


def create_iep_plan_table(document, data):
    """III-Bireyselleştirilmiş Eğitim Planı bölümünü oluşturur."""
    # Bu tablo çok karmaşık olduğu için genel stil yerine sadece başlıkları renklendiriyoruz.
    document.add_heading('III-Bireyselleştirilmiş Eğitim Planı', level=1)
    dersler = data.get("secilenDersler", [])
    if not dersler:
        document.add_paragraph("Öğrenci için seçilmiş ders veya hedef bulunmamaktadır.")
        return

    plan_table = document.add_table(rows=2, cols=9)
    plan_table.style = 'Table Grid'
    headers1 = ["Uzun Dönemli Amaçlar", "Kısa Dönemli Amaçlar", "Ölçüt", "Yöntem ve Teknik", "Kullanılacak Materyaller", "Başlama ve Bitiş Tarihi", "Ölçme ve Değerlendirme", "", ""]
    headers2 = ["", "", "", "", "", "", "Değerlendirme Yöntem ve Teknikleri", "Değerlendirme Tarihleri", "Performans"]

    for i, text in enumerate(headers1):
        set_cell_shading(plan_table.cell(0, i), '#4F81BD')
        plan_table.cell(0, i).text = text
        plan_table.cell(0, i).paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for i, text in enumerate(headers2):
        set_cell_shading(plan_table.cell(1, i), '#4F81BD')
        plan_table.cell(1, i).text = text
        plan_table.cell(1, i).paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for i in range(6): plan_table.cell(0, i).merge(plan_table.cell(1, i))
    plan_table.cell(0, 6).merge(plan_table.cell(0, 8))

    for ders in dersler:
        ders_row_cell = plan_table.add_row().cells[0]
        ders_row_cell.merge(plan_table.rows[-1].cells[-1])
        set_cell_shading(ders_row_cell, '#D3DFEE')
        set_cell_text(ders_row_cell, ders.get("dersAdi", "").upper(), bold=True, align_center=True)

        for uda in ders.get("uzunDonemliAmaclar", []):
            kda_list = uda.get("kisaDonemliAmaclar", [])
            if not kda_list: continue

            start_row_index = len(plan_table.rows)
            for kda in kda_list:
                kda_row_cells = plan_table.add_row().cells
                set_cell_shading(kda_row_cells[0], '#D3DFEE') # UDA sütunu
                kda_row_cells[1].text = kda.get("kdaMetni", "")
                kda_row_cells[2].text = kda.get("olcut", "")
                kda_row_cells[3].text = ", ".join(kda.get("ogretimYontemleri", []))
                kda_row_cells[4].text = ", ".join(kda.get("kullanilanMateryaller", []))
                kda_row_cells[5].text = f"{kda.get('baslamaTarihi', '')} - {kda.get('bitisTarihi', '')}"
                kda_row_cells[6].text = "Gözlem, Kontrol Listesi"
                kda_row_cells[7].text = kda.get("degerlendirmeTarihi", kda.get('bitisTarihi', ''))
                kda_row_cells[8].text = "" # Performans

            end_row_index = len(plan_table.rows) - 1
            b_cells = plan_table.rows[start_row_index].cells
            e_cells = plan_table.rows[end_row_index].cells
            b_cells[0].merge(e_cells[0])
            set_cell_text(b_cells[0], uda.get("udaMetni", ""), align_center=True)

def create_decisions_section(document, data):
    """IV-BEP Geliştirme Birim Kararları bölümünü oluşturur."""
    document.add_heading('IV-BEP Geliştirme Birimi Kararları', level=1)

    document.add_heading('B. Aile Bilgilendirme Süreci', level=2)
    aile_table = document.add_table(rows=1, cols=2)
    aile_table.cell(0, 0).text = "Aile Bilgilendirme Süreci"
    aile_table.cell(0, 0).merge(aile_table.cell(0, 1))

    aile_bilgi_data = [
        ("Aile öğrencinin gelişimi ile ilgili hangi sıklıkla bilgilendirilecek?", data.get("bilgilendirmeSikligi", "")),
        ("Aile öğrencinin gelişimi ile ilgili hangi yolla bilgilendirilecek?", ", ".join(data.get("bilgilendirmeYollari", []))),
        ("Aile eğitimi yapılacak m��?", "Evet" if data.get("aileEgitimiYapilacakMi", False) else "Hayır")
    ]
    for key, value in aile_bilgi_data:
        row = aile_table.add_row().cells
        row[0].text = key
        row[1].text = value

    apply_custom_table_style(aile_table)
    document.add_paragraph()

    document.add_heading('C. Diğer Kararlar', level=2)
    for karar in data.get("digerKararlar", []):
        document.add_paragraph(karar, style='List Bullet')
    document.add_paragraph(f"Bir Sonraki BEP geliştirme birimi toplantı tarihi: {data.get('sonrakiToplantiTarihi', '')}")

    document.add_paragraph()
    p_genel = document.add_paragraph()
    p_genel.add_run('Genel Bep Değerlendirmesi: ').bold = True
    p_genel.add_run("Bu plan hazırlanırken öğrencinin içinde bulunduğu durum ve ailevi koşullar göz önüne alınmış olup öğrenci için özel olarak bir yıllık plan şeklinde hazırlanmıştır. Gerekli görüldüğü takdirde bep birimiyle olağanüstü toplantı yaparak plan gözden geçirilip revize edilebilir.")


def create_members_table(document, data):
    """V-BEP Geliştirme Birimi Üyeleri bölümünü oluşturur."""
    document.add_heading('V-BEP Geliştirme Birimi Üyeleri', level=1)
    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Unvanı'
    hdr_cells[1].text = 'Adı Soyadı'
    hdr_cells[2].text = 'İmza'

    veli_secimi = data.get("veliSecimi", "")
    veli_ad_soyad = ""
    if veli_secimi == "Anne": veli_ad_soyad = data.get("anneAdSoyad", "")
    elif veli_secimi == "Baba": veli_ad_soyad = data.get("babaAdSoyad", "")
    else: veli_ad_soyad = data.get("digerVeliAdSoyad", "")

    uyeler = [
        ("Müdür/Müdür Yardımcısı", data.get("mudurAdi")),
        ("Sınıf Öğretmeni", data.get("sinifOgretmeni")),
        ("Veli", veli_ad_soyad),
        ("Rehber Öğretmen", data.get("rehberOgretmen")),
    ]
    for unvan, ad in uyeler:
        row_cells = table.add_row().cells
        row_cells[0].text = unvan
        row_cells[1].text = str(ad) if ad is not None else ''

    if data.get("alanOgretmenleri"):
        alan_ogretmen_row = table.add_row().cells
        alan_ogretmen_row[0].text = "Öğrencinin Dersini Okutan Alan Öğretmenleri"
        alan_ogretmen_row[0].merge(alan_ogretmen_row[2])
        for ogretmen in data.get("alanOgretmenleri", []):
            row_cells = table.add_row().cells
            row_cells[0].text = ogretmen.get("brans")
            row_cells[1].text = ogretmen.get("adSoyad")

    apply_custom_table_style(table)
    document.add_paragraph("\n\n")
    p_onay = document.add_paragraph()
    p_onay.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_onay.add_run(f"Uygundur\n{datetime.datetime.now().strftime('%d.%m.%Y')}\n")
    p_onay.add_run(f"{data.get('mudurAdi', 'Okul Müdürü Adı')}\n").bold = True
    p_onay.add_run("Okul Müdürü")


# --- ANA CLOUD FUNCTION (GÜNCELLENDİ) ---
@https_fn.on_request() # generate_bep_docx için memory ayarı yoksa varsayılan kullanılır
def generate_bep_docx(req: https_fn.Request) -> https_fn.Response:
    """Flutter'dan gelen BEP Plan�� verileriyle Word belgesi oluşturur."""
    try:
        data = req.get_json(silent=True)
        if not data:
            raise ValueError("Gelen veri boş.")

        ogrenci_ad_soyad = data.get('ogrenciAdSoyad', 'plan').replace(' ', '_')
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        unique_filename = f"BEP_{ogrenci_ad_soyad}_{timestamp}.docx"

        document = Document()

        # Dikey sayfalar
        create_cover_page(document, data)
        document.add_page_break()
        create_student_info_table(document, data)
        document.add_page_break()
        create_performance_form(document, data)

        # Yatay sayfa bölümü
        section_h = document.add_section(WD_ORIENT.LANDSCAPE)
        section_h.page_width, section_h.page_height = section_h.page_height, section_h.page_width

        create_iep_plan_table(document, data)

        # Tekrar dikey sayfa bölümü
        section_v = document.add_section(WD_ORIENT.PORTRAIT)
        # Genişlik ve yüksekliği orijinal dikey boyutlarına döndür
        # Genellikle ilk bölümün boyutları referans alınır.
        base_width = document.sections[0].page_width
        base_height = document.sections[0].page_height
        section_v.page_width = base_width
        section_v.page_height = base_height
        section_v.orientation = WD_ORIENT.PORTRAIT

        create_decisions_section(document, data)
        document.add_page_break()
        create_members_table(document, data)

        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)

        return https_fn.Response(
            buffer.getvalue(),
            status=200,
            headers={
                "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "Content-Disposition": f"attachment; filename=\"{unique_filename}\""
            }
        )
    except Exception as e:
        print(f"Hata oluştu: {e}")
        return https_fn.Response(json.dumps({"error": f"Sunucuda bir hata oluştu: {str(e)}"}), status=500, mimetype="application/json")

# YENİ: KABA DEĞERLENDİRME FORMU OLUŞTURMA HTTP FONKSİYONU
@https_fn.on_request(timeout_sec=300, memory=options.MemoryOption.MB_512) # DEĞİŞİKLİK: options.MemoryOption kullanıldı
def create_kaba_degerlendirme_word_http(req: https_fn.Request) -> https_fn.Response:
    try:
        # Flutter'dan gelen JSON verisini al
        data = req.get_json(silent=True)
        if not data:
            return https_fn.Response("Hatalı istek: JSON verisi bulunamadı.", status=400)

        # Kurum adı gibi Flutter'dan gelmesi beklenen ama Python'da varsayılanı olan alanlar için kontrol
        if 'kurum_adi' not in data or not data['kurum_adi']:
            data['kurum_adi'] = "Kurum Adı Belirtilmedi (Flutter)" # Flutter'dan gelmezse varsayılan

        # Word belgesini oluştur
        doc_bytes_io = generate_kaba_degerlendirme_docx(data)

        return https_fn.Response(
            doc_bytes_io.getvalue(),
            status=200,
            headers={
                "Content-Disposition": "attachment; filename=kaba_degerlendirme_formu.docx",
                "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            }
        )
    except Exception as e:
        print(f"Kaba Değerlendirme Word oluşturulurken hata oluştu: {e}")
        # İstemciye daha anlamlı bir hata mesajı göndermek için
        error_message = str(e)
        # Geliştirme aşamasında detaylı hata, canlıda daha genel bir mesaj olabilir
        return https_fn.Response(f"Sunucu hatası: Belge oluşturulamadı. Detay: {error_message}", status=500)


# ... (varsa diğer fonksiyonlarınız) ...

